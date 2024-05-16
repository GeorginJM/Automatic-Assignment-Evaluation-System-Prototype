from flask import Flask, request, jsonify, render_template,redirect, url_for, session, send_from_directory
from flask_mysqldb import MySQL
from werkzeug.utils import secure_filename
import os
import docx  # Import the docx module for handling .docx files
import tensorflow as tf
from transformers import BertTokenizer, TFBertForSequenceClassification
from datetime import datetime  # Import datetime module

app = Flask(__name__, template_folder='templates')

# MySQL Configuration
app.config['MYSQL_HOST'] = 'localhost'
app.config['MYSQL_USER'] = 'root'
app.config['MYSQL_PASSWORD'] = 'georgin'
app.config['MYSQL_DB'] = 'autoeval'

# File Upload Configuration
UPLOAD_FOLDER = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'uploads')
ALLOWED_EXTENSIONS = {'docx'}
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

app.secret_key = 'your_secret_key_here'

mysql = MySQL(app)

# Define custom loss function with L2 regularization and weighting scheme
def custom_loss(y_true, y_pred):
    # Cast y_true to float32
    y_true = tf.cast(y_true, tf.float32)
    
    # Define regularization strength
    l2_regularization = 0.001
    
    # Compute L2 regularization term
    regularization_term = tf.reduce_sum([tf.reduce_sum(tf.square(w)) for w in model.trainable_weights])
    regularization_loss = l2_regularization * regularization_term
    
    # Define weights for different classes (example)
    class_weights = tf.constant([1.0, 2.0])  # Adjust weights as needed
    
    # Compute weighted cross-entropy loss
    weighted_loss = tf.reduce_mean(tf.nn.weighted_cross_entropy_with_logits(labels=y_true, logits=y_pred, pos_weight=class_weights))
    
    # Compute total loss
    total_loss = weighted_loss + regularization_loss
    
    return total_loss


# Load BERT model for sequence classification
model = TFBertForSequenceClassification.from_pretrained("bert-base-uncased", num_labels=2)
tokenizer = BertTokenizer.from_pretrained("bert-base-uncased")

@app.route('/')
def index():
    # You can render the index.html file here
    return render_template('index.html')

@app.route('/authentication')
def authentication():
    # You can render the authentication.html file here
    return render_template('authentication.html')

@app.route('/teacherauthentication')
def teacherauthentication():
    # You can render the authentication.html file here
    return render_template('teacherauthentication.html')

@app.route('/frontpage')
def frontpage():
    # Fetch teacher emails
    cur = mysql.connection.cursor()
    cur.execute("SELECT username FROM teachers")
    teachers = cur.fetchall()
    cur.close()

    return render_template('frontpage.html', teachers=teachers)

@app.route('/frontpage1')
def frontpage1():
    return render_template('frontpage1.html')

@app.route('/assignments/<teacher_email>')
def teacher_assignments(teacher_email):
    # Fetch assignments for the given teacher email
    cur = mysql.connection.cursor()
    cur.execute("SELECT * FROM assignments WHERE teacher_id = (SELECT id FROM teachers WHERE username = %s)", (teacher_email,))
    assignments = cur.fetchall()
    cur.close()

    return render_template('teacher_assignments.html', assignments=assignments)

@app.route('/show_assignments')
def show_assignments():
    # Get the email of the currently logged-in teacher
    teacher_email = session.get('email')

    # Fetch assignments of the current teacher
    cur = mysql.connection.cursor()
    cur.execute("SELECT id, name, details FROM assignments WHERE teacher_id = (SELECT id FROM teachers WHERE username = %s)", (teacher_email,))
    assignments = cur.fetchall()
    cur.close()

    return render_template('show_assignments.html', assignments=assignments)

def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

@app.route('/post_assignment', methods=['GET', 'POST'])
def post_assignment():
    if request.method == 'POST':
        # Get assignment details from form data
        assignment_name = request.form.get('assignment_name')
        assignment_details = request.form.get('assignment_details')
        submission_deadline = request.form.get('last_submission_date')

        # Validate if submission_deadline is provided
        if submission_deadline:
            # Convert submission deadline to a datetime object
            submission_deadline = datetime.strptime(submission_deadline, '%Y-%m-%d')
        else:
            # If submission_deadline is not provided, set it to None or handle it accordingly
            submission_deadline = None

        # Get the teacher's email from the session
        teacher_email = session.get('email') if 'email' in session else None

        # Fetch the teacher_id corresponding to the teacher's email
        cur = mysql.connection.cursor()
        cur.execute("SELECT id FROM teachers WHERE username = %s", (teacher_email,))
        teacher_id = cur.fetchone()[0]
        cur.close()

        # Insert assignment details into the database along with submission deadline
        cur = mysql.connection.cursor()
        cur.execute("INSERT INTO assignments (name, details, last_submission_date, teacher_id) VALUES (%s, %s, %s, %s)", (assignment_name, assignment_details, submission_deadline, teacher_id))
        mysql.connection.commit()
        cur.close()

        return redirect(url_for('frontpage1'))

    return render_template('post_assignment.html')

CORRECTNESS_THRESHOLD=0.75

@app.route('/submit_assignment/<int:assignment_id>', methods=['GET', 'POST'])
def submit_assignment(assignment_id):
    last_submission_date = None  # Initialize with a default value

    # Get the currently logged-in user's email from the session
    email = session.get('email') if 'email' in session else None

    if request.method == 'POST':
        # Check if the assignment exists and get its details including the last submission date
        cur = mysql.connection.cursor()
        cur.execute("SELECT last_submission_date FROM assignments WHERE id = %s", (assignment_id,))
        assignment = cur.fetchone()
        cur.close()

        if assignment:
            last_submission_date = assignment[0]  # Access the first element of the tuple
            
            # Check if the current date is before the last submission date
            if datetime.now().date() <= last_submission_date:
                # Continue with submission process
                
                # Check if the post request has the file part
                if 'file' not in request.files:
                    return redirect(request.url)
                file = request.files['file']
                # If the user does not select a file, the browser submits an
                # empty file without a filename.
                if file.filename == '':
                    return redirect(request.url)
                if file and allowed_file(file.filename):
                    filename = secure_filename(file.filename)
                    file.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))

                    # Insert file location, email, and assignment ID into the database
                    cur = mysql.connection.cursor()
                    cur.execute("INSERT INTO submissions (assignment_id, email, file_location) VALUES (%s, %s, %s)", (assignment_id, email, filename))
                    mysql.connection.commit()
                    cur.close()

                    # Extract question-answer pairs from the uploaded .docx file
                    question_answer_pairs = extract_question_answer_pairs(os.path.join(app.config['UPLOAD_FOLDER'], filename))

                    # Load your TensorFlow model
                    loaded_model = tf.keras.models.load_model(r"C:\Users\Admin\Desktop\Georgin\Studies\autoevalsystem\z.others\savedmodel_latest(3)", custom_objects={'custom_loss': custom_loss})

                    # Process submissions and evaluate them
                    evaluation_results = evaluate_question_answer_pairs(question_answer_pairs, loaded_model)

                    # Store evaluation results in the database
                    if store_evaluation_results(assignment_id, evaluation_results,email):
                        print("Evaluation results stored successfully.")
                    else:
                        print("Failed to store evaluation results.")

                    return redirect(url_for('frontpage'))  # Redirect to the frontpage after submitting assignment
            else:
                return "You cannot submit the assignment after the deadline."
        else:
            return "Assignment not found."
    
    else:
        # Fetch last submission date from the database
        cur = mysql.connection.cursor()
        cur.execute("SELECT last_submission_date FROM assignments WHERE id = %s", (assignment_id,))
        assignment = cur.fetchone()
        cur.close()

        if assignment:
            last_submission_date = assignment[0]
    
    # Render the template for submitting the assignment
    return render_template('submit_assignment.html', assignment_id=assignment_id, last_submission_date=last_submission_date)

def extract_question_answer_pairs(file_path):
    # Initialize an empty list to store question-answer pairs
    question_answer_pairs = []

    # Open the .docx file
    doc = docx.Document(file_path)

    # Iterate through each paragraph in the document
    for i in range(len(doc.paragraphs)):
        # Check if the paragraph starts with "Question:"
        if doc.paragraphs[i].text.strip().startswith("Question:"):
            question = doc.paragraphs[i].text.strip().replace("Question:", "").strip()
            answer = ""

            # Accumulate answer until the next question or end of document
            while i+1 < len(doc.paragraphs) and not doc.paragraphs[i+1].text.strip().startswith("Question:"):
                i += 1
                answer += doc.paragraphs[i].text.strip() + "\n"

            # Append the question-answer pair to the list
            question_answer_pairs.append((question, answer.strip()))

    return question_answer_pairs

def evaluate_question_answer_pairs(question_answer_pairs, model):
    # Initialize an empty list to store evaluation results
    evaluation_results = []

    # Iterate through each question-answer pair
    for question, answer in question_answer_pairs:
        # Perform any preprocessing on the question and answer if needed

        # Tokenize the question and answer using the BERT tokenizer
        inputs = tokenizer(question, answer, padding=True, truncation=True, return_tensors="tf")

        # Perform inference using the machine learning model
        outputs = model(inputs)

        # Extract the predicted probability of correctness
        probability_correct = tf.nn.softmax(outputs['logits'], axis=-1).numpy()[0][1]

        # Check if the probability meets the correctness threshold
        if probability_correct >= CORRECTNESS_THRESHOLD:
            is_correct = True
        else:
            is_correct = False

        # Append the evaluation result to the list
        evaluation_results.append((question, answer, probability_correct, is_correct))

    return evaluation_results

def store_evaluation_results(assignment_id, evaluation_results,email):
    try:
        # Connect to the database
        conn = mysql.connection
        cursor = conn.cursor()

        # Iterate over evaluation results and insert them into the database
        for question, answer, probability_correct, is_correct in evaluation_results:
            # Prepare SQL statement to insert evaluation result into the database
            sql = "INSERT INTO evaluation_results (assignment_id, question, answer, probability_correct, is_correct,submission_email) VALUES (%s, %s, %s, %s, %s,%s)"
            values = (assignment_id, question, answer, probability_correct, is_correct ,email)

            # Execute SQL statement
            cursor.execute(sql, values)

        # Commit the transaction
        conn.commit()

        # Close the cursor
        cursor.close()

        return True  # Return True if storing evaluation results succeeds

    except Exception as e:
        # Rollback the transaction if an error occurs
        conn.rollback()
        cursor.close()
        print("Error storing evaluation results:", e)
        return False  # Return False if storing evaluation results fails

@app.route('/view_submissions/<int:assignment_id>')
def view_submissions(assignment_id):
    # Fetch submissions and evaluations from the database
    cur = mysql.connection.cursor()
    cur.execute("SELECT * FROM submissions WHERE assignment_id = %s", (assignment_id,))
    submissions = cur.fetchall()

    cur.execute("SELECT * FROM evaluation_results WHERE assignment_id = %s", (assignment_id,))
    evaluations = cur.fetchall()

    cur.close()

    # Pass submissions and evaluations to the template
    return render_template('view_submission.html', submissions=submissions, evaluations=evaluations)

@app.route('/uploads/<filename>')
def uploaded_file(filename):
    return send_from_directory(app.config['UPLOAD_FOLDER'], filename)

@app.route('/login', methods=['POST'])
def login():
    username = request.form.get('username')
    password = request.form.get('password')
    user_type = request.form.get('user_type')

    cur = mysql.connection.cursor()

    if user_type == 'teacher':
        cur.execute("SELECT * FROM teachers WHERE username = %s", (username,))
        frontpage_route = 'frontpage1'  # Route for teacher front page
    elif user_type == 'student':
        cur.execute("SELECT * FROM students WHERE username = %s", (username,))
        frontpage_route = 'frontpage'  # Route for student front page
    else:
        return redirect(url_for('authentication', error='Invalid user type'))

    user = cur.fetchone()
    cur.close()

    if user:
        if password == user[2]:
            session['email'] = user[1]  # Assuming email is in the first column
            return redirect(url_for(frontpage_route))
        else:
            return redirect(url_for('authentication', error='Wrong password'))
    else:
        return redirect(url_for('authentication', error="User doesn't exist"))

@app.route('/logout', methods=['GET', 'POST'])
def logout():
    # Clear the user session
    session.pop('username', None)
    return redirect(url_for('index'))  # Redirect to the login page

if __name__ == '__main__':
    app.run(debug=True)


