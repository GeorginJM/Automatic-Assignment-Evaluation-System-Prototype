import os
import tensorflow as tf
from transformers import BertTokenizer, TFBertForSequenceClassification
from sklearn.model_selection import train_test_split
import docx
import numpy as np
import random
from nltk.corpus import wordnet
import nltk

# Download NLTK resources
nltk.download('wordnet')

# Define custom loss function with L2 regularization and weighting scheme
def custom_loss(y_true, y_pred):
    # Cast y_true to float32
    y_true = tf.cast(y_true, tf.float32)

    # Define regularization strength
    l2_regularization = 0.001

    # Compute L2 regularization term
    regularization_term = tf.reduce_sum([tf.reduce_sum(tf.square(w)) for w in model.trainable_weights])
    regularization_loss = l2_regularization * regularization_term

    # Define weights for different classes (example)
    class_weights = tf.constant([1.0, 2.0])  # Adjust weights as needed

    # Compute weighted cross-entropy loss
    weighted_loss = tf.reduce_mean(tf.nn.weighted_cross_entropy_with_logits(labels=y_true, logits=y_pred, pos_weight=class_weights))

    # Compute total loss
    total_loss = weighted_loss + regularization_loss

    return total_loss

def read_question_answer_pairs(file_path):
    pairs = []
    doc = docx.Document(file_path)
    i = 0
    while i < len(doc.paragraphs):
        # Check if the paragraph starts with a number followed by a period
        if doc.paragraphs[i].text.strip().startswith(str(len(pairs) + 1) + "."):
            question = doc.paragraphs[i].text.strip()
            i += 1
            answer = ""
            # Accumulate answer until the next question or end of document
            while i < len(doc.paragraphs) and not doc.paragraphs[i].text.strip().startswith(str(len(pairs) + 2) + "."):
                answer += doc.paragraphs[i].text.strip() + "\n"
                i += 1
            pairs.append((question, answer.strip()))
        else:
            i += 1
    return pairs

def generate_adversarial_answer(original_answer):
    # Tokenize answer
    tokens = tokenizer.encode(original_answer, add_special_tokens=False)
    adversarial_tokens = tokens.copy()

    # Randomly delete some words
    adversarial_tokens = delete_words(adversarial_tokens)

    # Replace some words with antonyms
    adversarial_tokens = replace_with_antonyms(adversarial_tokens)

    # Insert additional words or phrases
    adversarial_tokens = insert_words(adversarial_tokens)

    # Introduce negations
    adversarial_tokens = introduce_negations(adversarial_tokens)

    return tokenizer.decode(adversarial_tokens, skip_special_tokens=True)

def delete_words(tokens):
    num_words_to_delete = random.randint(1, min(5, len(tokens) // 2))
    for _ in range(num_words_to_delete):
        index = random.randint(0, len(tokens) - 1)
        tokens.pop(index)
    return tokens

def replace_with_antonyms(tokens):
    for i, token_id in enumerate(tokens):
        word = tokenizer.decode(token_id)
        antonyms = get_antonyms(word)
        if antonyms:
            antonym = random.choice(antonyms)
            antonym_token_id = tokenizer.encode(antonym, add_special_tokens=False)[0]
            tokens[i] = antonym_token_id
    return tokens

def get_antonyms(word):
    antonyms = []
    for syn in wordnet.synsets(word):
        for lemma in syn.lemmas():
            if lemma.antonyms():
                antonyms.append(lemma.antonyms()[0].name())
    return antonyms

def insert_words(tokens):
    num_words_to_insert = random.randint(1, 3)
    for _ in range(num_words_to_insert):
        index = random.randint(0, len(tokens) - 1)
        random_word = random.choice(['nonsense', 'gibberish', 'confusion'])
        random_word_token_id = tokenizer.encode(random_word, add_special_tokens=False)[0]
        tokens.insert(index, random_word_token_id)
    return tokens

def introduce_negations(tokens):
    for i, token_id in enumerate(tokens):
        word = tokenizer.decode(token_id)
        if word in ['is', 'are', 'have', 'has', 'does']:
            negated_word = 'does not' if word == 'does' else word + ' not'
            negated_word_token_id = tokenizer.encode(negated_word, add_special_tokens=False)[0]
            tokens[i] = negated_word_token_id
    return tokens

# Load question-answer pairs from the .docx file
data_file_path = r"C:\Users\Admin\Desktop\Georgin\Studies\autoevalsystem\qanda.docx"
question_answer_pairs = read_question_answer_pairs(data_file_path)

# Tokenize data using BERT tokenizer
tokenizer = BertTokenizer.from_pretrained("bert-base-uncased")

# Generate adversarial answers
adversarial_answers = [generate_adversarial_answer(pair[1]) for pair in question_answer_pairs]

# Combine original answers and adversarial answers
all_answers = [pair[1] for pair in question_answer_pairs] + adversarial_answers
all_labels = np.concatenate((np.ones(len(question_answer_pairs)), np.zeros(len(adversarial_answers))))

# Tokenize data
encoded_data = tokenizer(all_answers, padding=True, truncation=True, return_tensors="tf")

# Prepare labels
labels = tf.one_hot(all_labels, depth=2)

# Split data into training and validation sets
train_inputs, val_inputs, train_labels, val_labels = train_test_split(
    encoded_data["input_ids"].numpy(),
    labels.numpy(),
    test_size=0.2,
    random_state=42
)

# Split attention masks in the same way
train_attention_masks, val_attention_masks, _, _ = train_test_split(
    encoded_data["attention_mask"].numpy(),
    labels.numpy(),
    test_size=0.2,
    random_state=42
)

# Load BERT model for sequence classification
model = TFBertForSequenceClassification.from_pretrained("bert-base-uncased", num_labels=2)

# Fine-tune only certain layers
for layer in model.layers[:-4]:
    layer.trainable = False

# Compile model with custom loss function
optimizer = tf.keras.optimizers.Adam(learning_rate=3e-5)
model.compile(optimizer=optimizer, loss=custom_loss, metrics=['accuracy'])

# Early Stopping callback
early_stopping = tf.keras.callbacks.EarlyStopping(monitor='val_loss', patience=3, restore_best_weights=True)

# Train model
epochs = 3
batch_size = 32
history = model.fit(
    [train_inputs, train_attention_masks],
    train_labels,
    validation_data=([val_inputs, val_attention_masks], val_labels),
    epochs=epochs,
    batch_size=batch_size,
    callbacks=[early_stopping],
    verbose=1
)

# Evaluate model
eval_results = model.evaluate([val_inputs, val_attention_masks], val_labels)
print("Evaluation results:", eval_results)

# Save the fine-tuned model
model.save(r"C:\Users\Admin\Desktop\Georgin\Studies\autoevalsystem\savedmodel_latest(4)")
